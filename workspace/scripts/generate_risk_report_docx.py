#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador DOCX para informes de riesgos SaaS.

Entrada principal:
  /data/cases/<case_id>/04_report/report_render_data.json

Plantilla:
  /data/templates/NuevaPlantillaInforme_MARKERS.docx

Salida:
  /data/cases/<case_id>/05_report/Informe_Ciberseguridad_<case_id>.docx

Notas:
- El Excel del cuestionario se referencia como evidencia visible con hipervínculo.
- La incrustación OLE real de Excel requiere Word/COM en Windows o librería comercial; no es portable en Docker/Linux.
"""

import argparse
import json
import re
import shutil
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

LEVEL_COLORS_HEX = {
    "ALTO": "FF0000",
    "MEDIO": "FFC000",
    "BAJO": "FFD966",
    "SIN RIESGO": "70AD47",
    "": "D9D9D9",
}

PIE_COLORS = {
    "BAJO": "#FFD966",
    "MEDIO": "#FFC000",
    "ALTO": "#FF0000",
}

LEVEL_RANK = {"": 0, "BAJO": 1, "MEDIO": 2, "ALTO": 3}

# Paleta visual EMPRESA / informe
ORANGE = "FF8200"
ORANGE_DARK = "FF6600"
TEAL = "008A9A"
TEAL_DARK = "007C89"
LIGHT_GRAY = "F2F2F2"
BORDER_GRAY = "D9D9D9"
WHITE = "FFFFFF"
BLACK = "000000"

# Colores para riesgo global y niveles
RISK_LEVEL_COLORS_HEX = {
    "ALTO": "FF0000",
    "MEDIO": "FF8200",
    "BAJO": "FFFF00",
    "SIN RIESGO": "70AD47",
    "": "BFBFBF",
}


def load_json(path):
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"No existe el JSON requerido: {path}")
    return json.loads(path.read_text(encoding='utf-8'))


def save_json(path, data):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')


def safe_text(value):
    if value is None:
        return ""
    if isinstance(value, (list, tuple)):
        return "\n".join(safe_text(v) for v in value if safe_text(v))
    return str(value).strip()


def clean_text(value):
    text = safe_text(value)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_level(value):
    v = safe_text(value).upper()
    if "ALTO" in v:
        return "ALTO"
    if "MEDIO" in v:
        return "MEDIO"
    if "BAJO" in v:
        return "BAJO"
    if "SIN RIESGO" in v:
        return "SIN RIESGO"
    return ""


def mitigation_time(level):
    level = normalize_level(level)
    if level == "ALTO":
        return "3 meses"
    if level == "MEDIO":
        return "12 meses"
    if level == "BAJO":
        return "Pte."
    return ""


def clean_cover_value(value):
    """Limpia valores de portada y evita los ':' o puntos finales que ya añade la plantilla."""
    text = clean_text(value)
    text = text.rstrip(" .:;")
    return text


def is_weak_project_value(value):
    text = clean_cover_value(value).lower()
    return text in {"", "no", "si", "sí", "n/a", "na", "none", "null", "undefined", "pendiente de confirmar"}


def get_nested(data, path, default=None):
    cur = data
    for part in path:
        if isinstance(cur, dict) and part in cur:
            cur = cur[part]
        else:
            return default
    return cur


def first_non_empty(*values):
    for value in values:
        text = clean_text(value)
        if text and text.lower() not in {"none", "null", "undefined"}:
            return text
    return ""


def get_triaje(data):
    return (
        data.get("triaje")
        or get_nested(data, ["project_context_input", "triaje"], {})
        or {}
    )


def get_project_name(data):
    triaje = get_triaje(data)
    project_code = get_nested(triaje, ["contract", "project_code"], "")
    candidates = [
        project_code,
        get_nested(data, ["project_context_input", "proyecto"], ""),
        get_nested(data, ["cover", "proyecto"], ""),
        get_nested(data, ["executive_visual", "nombre_proyecto"], ""),
        data.get("case_id", ""),
    ]
    for c in candidates:
        value = clean_cover_value(c)
        if value and not is_weak_project_value(value):
            return value
    return clean_cover_value(data.get("case_id", "Proyecto"))


def get_business_line(data):
    return clean_cover_value(first_non_empty(
        get_nested(data, ["project_context_input", "linea_negocio"], ""),
        get_nested(data, ["cover", "negocio"], ""),
        get_nested(data, ["executive_visual", "linea_negocio"], ""),
        "Pendiente de confirmar"
    ))


def get_provider_name(data):
    triaje = get_triaje(data)
    providers = get_nested(triaje, ["contract", "providers"], [])
    provider_from_triaje = ", ".join(clean_text(p) for p in providers if clean_text(p)) if isinstance(providers, list) else clean_text(providers)
    return clean_cover_value(first_non_empty(
        get_nested(data, ["project_context_input", "proveedor"], ""),
        get_nested(data, ["cover", "proveedor"], ""),
        provider_from_triaje,
        "Proveedor"
    ))


def get_platform_name(data):
    triaje = get_triaje(data)
    return clean_cover_value(first_non_empty(
        get_nested(triaje, ["solution", "saas_name"], ""),
        get_nested(data, ["project_context_input", "nombre_solucion"], ""),
        get_nested(data, ["project_context_input", "saas_name"], ""),
        get_project_name(data)
    ))


def risk_color_hex(level):
    return RISK_LEVEL_COLORS_HEX.get(normalize_level(level), "BFBFBF")


def classification_with_gdpr(data):
    triaje = get_triaje(data)
    information = get_nested(triaje, ["information"], {}) or {}

    classification = first_non_empty(
        information.get("classification"),
        get_nested(data, ["project_context_input", "tipo_informacion"], ""),
        get_nested(data, ["executive_visual", "clasificacion_informacion"], ""),
        "Pendiente de confirmar"
    )

    personal_raw = information.get("personal_data")
    if personal_raw in [None, ""]:
        personal_raw = get_nested(data, ["project_context_input", "datos_personales"], "")
    has_personal_data = personal_raw is True or str(personal_raw).strip().lower() in {"true", "si", "sí", "yes", "1"}

    regime = first_non_empty(
        information.get("personal_data_regime"),
        get_nested(data, ["project_context_input", "regimen_datos_personales"], "")
    ).upper()

    classification = clean_cover_value(classification)
    if has_personal_data and "GDPR" in regime and "GDPR" not in classification.upper():
        return f"{classification} + GDPR"
    return classification


def set_run_color_hex(run, hex_color):
    hex_color = str(hex_color or "000000").replace("#", "")
    run.font.color.rgb = RGBColor.from_string(hex_color)


def clear_cell(cell):
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    for run in p.runs:
        run.text = ""
    return p


def set_cell_border(cell, **kwargs):
    """Aplica bordes a una celda. Ejemplo: top={'val':'single','sz':'8','color':'FF8200'}"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs.get(edge) or {}
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, color=BORDER_GRAY, size="4"):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": size, "space": "0", "color": color},
                bottom={"val": "single", "sz": size, "space": "0", "color": color},
                left={"val": "single", "sz": size, "space": "0", "color": color},
                right={"val": "single", "sz": size, "space": "0", "color": color},
            )


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_dxa))
    tc_w.set(qn('w:type'), 'dxa')


def replace_marker_text_formatted(doc, marker, value, bold=False, color_hex=None, font_size=None, align=None):
    value = safe_text(value)
    replaced = False
    for p in iter_paragraphs_in_document(doc):
        if marker in p.text:
            full = p.text.replace(marker, value)
            for run in p.runs:
                run.text = ""
            run = p.runs[0] if p.runs else p.add_run()
            run.text = full
            run.bold = bold
            if color_hex:
                set_run_color_hex(run, color_hex)
            if font_size:
                run.font.size = Pt(font_size)
            if align == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            replaced = True
    return replaced


def get_heading2_style(doc):
    """Devuelve el estilo equivalente a Título 2.

    En python-docx, los estilos integrados suelen exponerse con el nombre
    inglés (Heading 2), aunque en Word en español se vean como Título 2.
    También se comprueba el style_id habitual de plantillas españolas.
    """
    candidate_names = [
        "Heading 2",
        "Título 2",
        "Titulo 2",
    ]

    for name in candidate_names:
        try:
            return doc.styles[name]
        except Exception:
            pass

    for style in doc.styles:
        if getattr(style, "style_id", "") in {"Ttulo2", "Heading2"}:
            return style

    return None


def clean_heading_title(value):
    """Elimina numeración manual para que Word la genere con el estilo Título 2."""
    text = safe_text(value).strip()
    # Si llega "7.1 Robo de información" o "7.1.1 Robo...", deja solo el texto.
    text = re.sub(r"^\s*\d+(?:\.\d+)*\s+", "", text)
    return text or "Riesgo pendiente"


def add_heading2_after_element(doc, ref_el, title):
    """Inserta un título de nivel 2 usando el estilo Título 2 / Heading 2.

    La numeración 7.1, 7.2, etc. debe generarla Word mediante la lista
    multinivel asociada al estilo de la plantilla. Por eso aquí NO se escribe
    numeración manual y NO se elimina la numeración del párrafo.
    """
    p = insert_paragraph_after_element(doc, ref_el, "")

    heading2_style = get_heading2_style(doc)
    if heading2_style is not None:
        p.style = heading2_style
    else:
        p.style = "Normal"

    p.add_run(clean_heading_title(title))
    return p

def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill_hex)


def set_cell_text(cell, text, bold=False, font_size=None, align='left', color_hex=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(safe_text(text))
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet_lines(cell, values, font_size=8.5):
    cell.text = ""
    values = [clean_text(v) for v in values if clean_text(v)]
    if not values:
        p = cell.paragraphs[0]
        run = p.add_run("Pendiente de confirmar")
        run.font.size = Pt(font_size)
        return
    for idx, value in enumerate(values):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        if len(values) > 1:
            try:
                p.style = 'List Bullet' if 'List Bullet' in [s.name for s in cell._parent._parent.part.document.styles] else p.style
            except Exception:
                pass
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        run.font.size = Pt(font_size)


def unique(values):
    out = []
    seen = set()
    for value in values or []:
        if isinstance(value, list):
            candidates = value
        else:
            candidates = [value]
        for c in candidates:
            text = clean_text(c)
            if not text:
                continue
            key = text.lower()
            if key not in seen:
                seen.add(key)
                out.append(text)
    return out


def iter_paragraphs_in_document(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_marker_text(doc, marker, value):
    value = safe_text(value)
    replaced = False
    for p in iter_paragraphs_in_document(doc):
        if marker in p.text:
            full = p.text.replace(marker, value)
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = full
            else:
                p.add_run(full)
            replaced = True
    return replaced



def replace_all_text_markers_in_xml(docx_path, replacements):
    """Sustituye markers que están dentro de shapes/textboxes.

    python-docx no ve los párrafos de los cuadros de texto, por eso hacemos una
    sustitución OOXML directa. Importante: esta función NO usa regex expansivas
    entre runs, porque eso puede romper el XML del documento si el marcador está
    dentro de un textbox. Solo sustituye texto literal y escapado XML.
    """
    from zipfile import ZipFile, ZIP_DEFLATED
    import tempfile
    from xml.sax.saxutils import escape

    docx_path = Path(docx_path)
    tmp_out = docx_path.with_suffix('.patched.docx')

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with ZipFile(docx_path) as zin:
            zin.extractall(td)

        for xml_path in (td / 'word').rglob('*.xml'):
            xml = xml_path.read_text(encoding='utf-8', errors='ignore')
            original = xml

            for marker, value in replacements.items():
                xml_value = escape(safe_text(value), entities={'"': '&quot;', "'": '&apos;'})
                xml = xml.replace(marker, xml_value)

            if xml != original:
                xml_path.write_text(xml, encoding='utf-8')

        with ZipFile(tmp_out, 'w', ZIP_DEFLATED) as zout:
            for path in td.rglob('*'):
                if path.is_file():
                    zout.write(path, path.relative_to(td))

    tmp_out.replace(docx_path)

def find_paragraph_with_marker(doc, marker):
    for p in iter_paragraphs_in_document(doc):
        if marker in p.text:
            return p
    return None


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""


def remove_paragraph(paragraph):
    el = paragraph._p
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_after(reference_p, text="", style=None):
    new_p = OxmlElement('w:p')
    reference_p._p.addnext(new_p)
    p = Paragraph(new_p, reference_p._parent)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if text:
        p.add_run(text)
    return p


def insert_table_after_paragraph(doc, reference_p, rows, cols, style='Table Grid'):
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = style
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    reference_p._p.addnext(tbl)
    return table




def insert_paragraph_after_element(doc, ref_el, text="", style=None):
    new_p = OxmlElement('w:p')
    ref_el.addnext(new_p)
    p = Paragraph(new_p, doc._body)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    if text:
        p.add_run(text)
    return p


def insert_table_after_element(doc, ref_el, rows, cols, style='Table Grid'):
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = style
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    ref_el.addnext(tbl)
    return table

def insert_picture_after_paragraph(doc, reference_p, image_path, width_inches=5.8):
    p = insert_paragraph_after(reference_p, "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    r_pr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink










def remove_ole_objects_inside_textboxes(docx_path):
    """Sanea el DOCX para evitar el aviso de Word:
    "No puede poner un objeto de dibujo dentro de un cuadro de texto...".

    En esta plantilla hay objetos Excel/OLE y dibujos dentro de cuadros de texto.
    Word puede abrir el DOCX, pero muestra un aviso de reparación/compatibilidad.
    Esta función elimina de forma estructural esos elementos conflictivos usando
    lxml, que ya está disponible porque python-docx depende de lxml.

    Limpieza aplicada:
    - elimina bloques mc:AlternateContent que contienen OLE/Package dentro de un
      cuadro de texto;
    - elimina párrafos dentro de w:txbxContent que contienen dibujos/OLE;
    - elimina objetos OLE globales;
    - elimina relaciones a paquetes embebidos y la carpeta word/embeddings.
    """
    from zipfile import ZipFile, ZIP_DEFLATED
    import tempfile
    import re
    from lxml import etree

    docx_path = Path(docx_path)
    tmp_out = docx_path.with_suffix('.sanitized_textboxes.docx')

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'o': 'urn:schemas-microsoft-com:office:office',
        'v': 'urn:schemas-microsoft-com:vml',
    }

    # Elementos que no deben aparecer dentro de un w:txbxContent.
    bad_inside_textbox_xpath = (
        './/w:object | .//o:OLEObject | .//w:drawing | .//w:pict | '
        './/mc:AlternateContent | .//v:shape | .//v:group'
    )

    # Indicadores de objeto incrustado/OLE.
    embedded_object_xpath = './/w:object | .//o:OLEObject'

    def remove_element(el):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            return True
        return False

    def parse_xml_bytes(raw):
        parser = etree.XMLParser(remove_blank_text=False, recover=True, huge_tree=True)
        return etree.fromstring(raw, parser=parser)

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with ZipFile(docx_path) as zin:
            zin.extractall(td)

        changed = False
        word_dir = td / 'word'

        # 1) XML del documento, headers, footers, footnotes, endnotes, etc.
        for xml_path in [p for p in word_dir.rglob('*.xml') if p.is_file()]:
            raw = xml_path.read_bytes()
            if b'txbxContent' not in raw and b'OLEObject' not in raw and b'<w:object' not in raw:
                continue

            try:
                root = parse_xml_bytes(raw)
            except Exception:
                # Fallback mínimo por regex si algún XML no se puede parsear.
                xml = raw.decode('utf-8', errors='ignore')
                original = xml
                xml = re.sub(r'<w:object\b[\s\S]*?</w:object>', '', xml, flags=re.S)
                xml = re.sub(r'<o:OLEObject\b[^>]*/>', '', xml, flags=re.S)
                if xml != original:
                    xml_path.write_text(xml, encoding='utf-8')
                    changed = True
                continue

            local_changed = False

            # 1.a) Si un AlternateContent contiene un textbox con OLE, se elimina
            # entero. Esto quita el icono/objeto incrustado problemático.
            for ac in list(root.xpath('.//mc:AlternateContent', namespaces=ns)):
                has_textbox = bool(ac.xpath('.//w:txbxContent', namespaces=ns))
                has_embedded_object = bool(ac.xpath(embedded_object_xpath, namespaces=ns))
                ac_xml = etree.tostring(ac, encoding='unicode')
                has_package_reference = 'relationships/package' in ac_xml or 'Target="embeddings/' in ac_xml or "Target='embeddings/" in ac_xml
                if has_textbox and (has_embedded_object or has_package_reference):
                    local_changed = remove_element(ac) or local_changed

            # 1.b) Dentro de cualquier textbox que sobreviva, eliminar párrafos
            # que contengan dibujos, VML, OLE o AlternateContent. Es la parte que
            # evita estrictamente "drawing object inside textbox".
            for txbx in list(root.xpath('.//w:txbxContent', namespaces=ns)):
                for p in list(txbx.xpath('./w:p', namespaces=ns)):
                    if p.xpath(bad_inside_textbox_xpath, namespaces=ns):
                        local_changed = remove_element(p) or local_changed

                # Si el textbox queda vacío, dejamos un párrafo válido mínimo.
                if not txbx.xpath('./w:p', namespaces=ns):
                    new_p = etree.Element('{%s}p' % ns['w'])
                    txbx.append(new_p)
                    local_changed = True

            # 1.c) Limpieza global de objetos OLE que puedan quedar fuera.
            for obj in list(root.xpath('.//w:object | .//o:OLEObject', namespaces=ns)):
                local_changed = remove_element(obj) or local_changed

            if local_changed:
                xml_path.write_bytes(etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=False))
                changed = True

        # 2) Eliminar relaciones a paquetes embebidos.
        for rel_path in [p for p in word_dir.rglob('*.rels') if p.is_file()]:
            xml = rel_path.read_text(encoding='utf-8', errors='ignore')
            original = xml
            xml = re.sub(
                r'<Relationship\b[^>]*(?:officeDocument/2006/relationships/package|Target="embeddings/|Target=\'embeddings/)[^>]*/>',
                '',
                xml,
                flags=re.S,
            )
            if xml != original:
                rel_path.write_text(xml, encoding='utf-8')
                changed = True

        # 3) Eliminar overrides de Content_Types relacionados con embeddings.
        content_types = td / '[Content_Types].xml'
        if content_types.exists():
            xml = content_types.read_text(encoding='utf-8', errors='ignore')
            original = xml
            xml = re.sub(
                r'<Override\b[^>]*PartName="/word/embeddings/[^>]*/>',
                '',
                xml,
                flags=re.S,
            )
            if xml != original:
                content_types.write_text(xml, encoding='utf-8')
                changed = True

        # 4) No incluir embeddings en el DOCX reconstruido.
        with ZipFile(tmp_out, 'w', ZIP_DEFLATED) as zout:
            for path in td.rglob('*'):
                if not path.is_file():
                    continue
                rel = path.relative_to(td)
                if str(rel).replace('\\', '/').startswith('word/embeddings/'):
                    changed = True
                    continue
                zout.write(path, rel)

    tmp_out.replace(docx_path)
    return changed

def set_update_fields_on_open(docx_path):
    """Indica a Word que actualice índice/TOC y campos al abrir el documento."""
    from zipfile import ZipFile, ZIP_DEFLATED
    import tempfile

    docx_path = Path(docx_path)
    tmp_out = docx_path.with_suffix('.fields.docx')

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with ZipFile(docx_path) as zin:
            zin.extractall(td)

        settings_path = td / 'word' / 'settings.xml'
        if settings_path.exists():
            xml = settings_path.read_text(encoding='utf-8', errors='ignore')
            if '<w:updateFields' not in xml:
                xml = xml.replace('</w:settings>', '<w:updateFields w:val="true"/></w:settings>')
                settings_path.write_text(xml, encoding='utf-8')

        with ZipFile(tmp_out, 'w', ZIP_DEFLATED) as zout:
            for path in td.rglob('*'):
                if path.is_file():
                    zout.write(path, path.relative_to(td))
    tmp_out.replace(docx_path)


def postprocess_dynamic_report_xml(docx_path, data):
    """Ajustes OOXML no accesibles cómodamente con python-docx.

    - Cambia el color del contorno punteado del recuadro del riesgo global.
    - Cambia, si aparece, el relleno del recuadro de riesgo global.
    """
    from zipfile import ZipFile, ZIP_DEFLATED
    import tempfile

    level = normalize_level(
        get_nested(data, ["global_risk", "level"], "")
        or get_nested(data, ["executive_visual", "riesgo_global"], "")
    )
    color = risk_color_hex(level)

    docx_path = Path(docx_path)
    tmp_out = docx_path.with_suffix('.dynamic.docx')

    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with ZipFile(docx_path) as zin:
            zin.extractall(td)

        for xml_path in (td / 'word').glob('*.xml'):
            xml = xml_path.read_text(encoding='utf-8', errors='ignore')
            original = xml

            idx = xml.find('name="Rectangle 6"')
            if idx != -1:
                start = max(0, idx - 1500)
                end = min(len(xml), idx + 3500)
                segment = xml[start:end]
                segment = re.sub(
                    r'(<a:ln\b[^>]*>.*?<a:solidFill>\s*<a:srgbClr val=")[0-9A-Fa-f]{6}("/>)',
                    rf'\g<1>{color}\2',
                    segment,
                    count=1,
                    flags=re.S
                )
                xml = xml[:start] + segment + xml[end:]

            label = safe_text(get_nested(data, ["global_risk", "label"], f"RIESGO {level}"))
            if label:
                pos = xml.find(label)
                if pos != -1:
                    start = max(0, pos - 900)
                    end = min(len(xml), pos + 500)
                    segment = xml[start:end]
                    segment = re.sub(
                        r'(<w:shd\b[^>]*w:fill=")[0-9A-Fa-f]{6}("[^>]*/>)',
                        rf'\g<1>{color}\2',
                        segment,
                        count=1
                    )
                    xml = xml[:start] + segment + xml[end:]

            if xml != original:
                xml_path.write_text(xml, encoding='utf-8')

        with ZipFile(tmp_out, 'w', ZIP_DEFLATED) as zout:
            for path in td.rglob('*'):
                if path.is_file():
                    zout.write(path, path.relative_to(td))
    tmp_out.replace(docx_path)


def make_risk_pie_chart(risk_totals, output_path):
    labels = ['Bajo', 'Medio', 'Alto']
    values = [int(risk_totals.get('BAJO', 0)), int(risk_totals.get('MEDIO', 0)), int(risk_totals.get('ALTO', 0))]
    colors = [PIE_COLORS['BAJO'], PIE_COLORS['MEDIO'], PIE_COLORS['ALTO']]

    fig = plt.figure(figsize=(7.2, 3.0), dpi=150)
    ax = fig.add_subplot(111)

    total = sum(values)
    if total == 0:
        values = [1]
        labels = ['Sin riesgos']
        colors = ['#D9D9D9']

    def autopct_func(pct):
        if total == 0:
            return ''
        val = int(round(pct * total / 100.0))
        return str(val) if val > 0 else ''

    wedges, texts, autotexts = ax.pie(
        values,
        colors=colors,
        startangle=90,
        counterclock=False,
        autopct=autopct_func,
        pctdistance=1.22,
        wedgeprops={'linewidth': 1, 'edgecolor': 'white'}
    )
    ax.set_title('Riesgos totales', fontsize=13, fontweight='bold')
    ax.legend(wedges, labels, loc='center left', bbox_to_anchor=(1.05, 0.5), frameon=False)
    ax.axis('equal')
    fig.tight_layout()
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(output_path, transparent=False, bbox_inches='tight')
    plt.close(fig)
    return output_path




def find_table_row_for_paragraph(doc, paragraph):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p._p is paragraph._p:
                        return table, row, cell
    return None, None, None


def insert_risk_chart_at_marker(doc, marker_p, image_path, width_inches=4.8):
    """Inserta el gráfico de riesgos y, si el marker está en la fila 'Riesgos totales', elimina visualmente la columna izquierda."""
    table, row, cell = find_table_row_for_paragraph(doc, marker_p)

    if table is not None and row is not None and len(row.cells) > 1:
        try:
            merged = row.cells[0]
            for c in row.cells[1:]:
                merged = merged.merge(c)
            set_cell_shading(merged, LIGHT_GRAY)
            p = clear_cell(merged)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(width_inches))
            return
        except Exception:
            pass

    clear_paragraph(marker_p)
    marker_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = marker_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def create_incumplimientos_table(doc, marker_p, risk_groups):
    headers = [
        'ID',
        'Incumplimientos\nidentificados',
        'Controles Asociados',
        'Amenaza',
        'Nivel de\nIncumplimiento',
        'Área de\nafectación\ndel riesgo'
    ]
    table = insert_table_after_paragraph(doc, marker_p, rows=2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, align='center', color_hex=WHITE, font_size=10)
        set_cell_shading(hdr[i], ORANGE)

    legend = table.rows[1].cells
    try:
        legend[0].merge(legend[-1])
        legend_cell = legend[0]
    except Exception:
        legend_cell = legend[0]
    set_cell_shading(legend_cell, WHITE)
    p = clear_cell(legend_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, color in [('Nivel de riesgo asociado alto', 'C00000'), ('Nivel de riesgo asociado medio', ORANGE), ('Nivel de riesgo asociado bajo', 'FFFF00')]:
        r = p.add_run(' ● ')
        r.font.size = Pt(18)
        set_run_color_hex(r, color)
        txt = p.add_run(f' {label}    ')
        txt.font.size = Pt(10)

    for group in risk_groups:
        row = table.add_row().cells
        add_bullet_lines(row[0], group.get('control_ids', []))
        add_bullet_lines(row[1], group.get('incumplimientos', []))
        add_bullet_lines(row[2], group.get('controles_asociados', []))
        add_bullet_lines(row[3], group.get('amenazas', []))
        level = normalize_level(group.get('level'))

        p = clear_cell(row[4])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dot = p.add_run('●')
        dot.font.size = Pt(20)
        set_run_color_hex(dot, risk_color_hex(level))
        txt = p.add_run(f'\n{level}')
        txt.bold = True

        set_cell_text(row[5], group.get('area_afectacion', 'Local al proyecto'), bold=True, align='center')
        for c in row:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    widths = [900, 2200, 2700, 2300, 1600, 1600]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                set_cell_width(cell, widths[i])

    set_table_borders(table, ORANGE, size="6")
    remove_paragraph(marker_p)
    return table


def create_results_sections(doc, marker_p, risk_groups):
    ref_el = marker_p._p
    for idx, group in enumerate(risk_groups, start=1):
        # El texto del título NO incluye numeración manual.
        # Word genera 7.1, 7.2, etc. al aplicar el estilo Título 2 / Heading 2.
        heading_text = (
            group.get('risk_key')
            or group.get('risk_name')
            or 'Riesgo pendiente'
        )
        title = add_heading2_after_element(doc, ref_el, heading_text)
        ref_el = title._p

        table = insert_table_after_element(doc, ref_el, rows=7, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ref_el = table._tbl

        set_cell_text(table.cell(0,0), group.get('report_code', f'R_{idx:03d}'), bold=True, align='center', color_hex=WHITE, font_size=11)
        set_cell_text(table.cell(0,1), group.get('risk_name', '').upper(), bold=True, align='center', color_hex=WHITE, font_size=11)
        set_cell_shading(table.cell(0,0), ORANGE)
        set_cell_shading(table.cell(0,1), ORANGE)

        level = normalize_level(group.get('level'))
        set_cell_text(table.cell(1,0), 'Nivel de riesgo', bold=True, color_hex=WHITE, font_size=10)
        set_cell_shading(table.cell(1,0), TEAL)
        set_cell_text(table.cell(1,1), level, bold=True, align='center', color_hex=risk_color_hex(level), font_size=13)
        set_cell_shading(table.cell(1,1), LIGHT_GRAY)

        set_cell_text(table.cell(2,0), 'Contexto', bold=True, color_hex=WHITE, font_size=10)
        set_cell_shading(table.cell(2,0), TEAL)
        add_bullet_lines(table.cell(2,1), group.get('contextos', []))
        set_cell_shading(table.cell(2,1), LIGHT_GRAY)

        set_cell_text(table.cell(3,0), 'Detalle\nIncumplimiento', bold=True, color_hex=WHITE, font_size=10)
        set_cell_shading(table.cell(3,0), TEAL)
        add_bullet_lines(table.cell(3,1), group.get('detalles', []))

        set_cell_text(table.cell(4,0), 'Medidas\nmitigantes', bold=True, color_hex=WHITE, font_size=10)
        set_cell_shading(table.cell(4,0), TEAL)
        add_bullet_lines(table.cell(4,1), group.get('medidas_mitigantes', []))

        set_cell_text(table.cell(5,0), 'Tiempo máximo\nde mitigación', bold=True, color_hex=WHITE, font_size=10)
        set_cell_shading(table.cell(5,0), TEAL)
        set_cell_text(table.cell(5,1), group.get('mitigation_time') or mitigation_time(level), bold=True, align='center')

        set_cell_text(table.cell(6,0), 'Nivel residual\ntras aplicar medidas', bold=True, color_hex=WHITE, font_size=10)
        set_cell_shading(table.cell(6,0), TEAL)
        set_cell_text(table.cell(6,1), 'SIN RIESGO', bold=True, align='center', color_hex='70AD47', font_size=11)
        set_cell_shading(table.cell(6,1), LIGHT_GRAY)

        for row in table.rows:
            set_cell_width(row.cells[0], 1900)
            set_cell_width(row.cells[1], 7400)
            for c in row.cells:
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_table_borders(table, BORDER_GRAY, size="4")

    remove_paragraph(marker_p)


def create_debt_table(doc, marker_p, deuda):
    table = insert_table_after_paragraph(doc, marker_p, rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_cell_text(table.rows[0].cells[0], 'MEDIDAS MITIGANTES', bold=True, align='center', color_hex=TEAL, font_size=12)
    set_cell_text(table.rows[0].cells[1], 'Tiempo\nmáximo de\nmitigación', bold=True, align='center', color_hex=TEAL, font_size=11)
    set_cell_shading(table.rows[0].cells[0], WHITE)
    set_cell_shading(table.rows[0].cells[1], WHITE)

    if not deuda:
        row = table.add_row().cells
        set_cell_text(row[0], 'No se han identificado medidas mitigantes pendientes.', align='center')
        set_cell_text(row[1], '')

    for item in deuda:
        row = table.add_row().cells
        level = normalize_level(item.get('level'))
        set_cell_text(row[0], item.get('medida_mitigante', 'Pendiente de confirmar'), align='center')
        set_cell_text(row[1], item.get('tiempo_maximo_mitigacion', mitigation_time(level)), align='center')
        set_cell_shading(row[0], LIGHT_GRAY)
        set_cell_shading(row[1], WHITE)

    for row in table.rows:
        set_cell_width(row.cells[0], 7200)
        set_cell_width(row.cells[1], 2000)
        for c in row.cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table, BORDER_GRAY, size="4")
    remove_paragraph(marker_p)


def insert_excel_reference(doc, marker_p, render_data, output_dir):
    src = safe_text(render_data.get('sources', {}).get('questionnaire_file'))
    evidence_dir = output_dir / 'evidencias'
    evidence_dir.mkdir(parents=True, exist_ok=True)

    p = marker_p
    clear_paragraph(p)

    if src and Path(src).exists():
        src_path = Path(src)
        target = evidence_dir / src_path.name
        try:
            shutil.copy2(src_path, target)
        except Exception:
            target = src_path
        p.add_run('Cuestionario SaaS cumplimentado: ').bold = True
        add_hyperlink(p, target.name, str(target))
    else:
        p.add_run('Cuestionario SaaS cumplimentado: ').bold = True
        p.add_run('Pendiente de adjuntar fichero Excel de cuestionario.')


def patch_basic_markers(doc, data):
    narrative = data.get('narrative', {}) or {}
    cv = data.get('control_versions', {}) or {}
    gv = data.get('global_risk', {}) or {}
    ev = data.get('executive_visual', {}) or {}
    chars = narrative.get('caracteristicas_generales', {}) or {}

    project_name = get_project_name(data)
    business_line = get_business_line(data)
    provider_name = get_provider_name(data)
    platform_name = get_platform_name(data)
    platform_provider_title = f"{platform_name} – {provider_name}".strip(" –")

    global_level = normalize_level(ev.get('riesgo_global') or gv.get('level'))
    global_label = clean_cover_value(gv.get('label') or f"RIESGO {global_level}".strip())

    replacements = {
        '{{PROYECTO}}': project_name,
        '{{NEGOCIO}}': business_line,
        '{{FECHA_VERSION}}': cv.get('fecha', datetime.now().strftime('%d/%m/%Y')),
        '{{RIESGO_GLOBAL_LABEL}}': global_label,
        '{{CONTEXTO_PROYECTO}}': narrative.get('contexto_proyecto', 'Pendiente de confirmar.'),
        '{{CARACTERISTICAS_ARQUITECTURA}}': chars.get('arquitectura', 'Información no aportada.'),
        '{{CARACTERISTICAS_AUTENTICACION}}': chars.get('autenticacion', 'Información no aportada.'),
        '{{CARACTERISTICAS_MONITORIZACION}}': chars.get('monitorizacion', 'Información no aportada.'),
        '{{CARACTERISTICAS_SEGURIDAD_SISTEMAS}}': chars.get('seguridad_sistemas', 'Información no aportada.'),
        '{{CARACTERISTICAS_CONTINUIDAD}}': chars.get('continuidad', 'Información no aportada.'),
        '{{RESUMEN_NOMBRE_PROYECTO}}': project_name,
        '{{RESUMEN_CLASIFICACION_INFORMACION}}': classification_with_gdpr(data),
        '{{RESUMEN_LINEA_NEGOCIO}}': business_line,
        '{{RESUMEN_NIVEL_AFECTACION}}': ev.get('nivel_afectacion', 'Local al proyecto'),
        '{{RESUMEN_RIESGO_GLOBAL}}': global_level,
        'Plataforma – Proveedor': platform_provider_title,
        'Plataforma - Proveedor': platform_provider_title,
    }

    special = {'{{RESUMEN_RIESGO_GLOBAL}}', '{{RESUMEN_NOMBRE_PROYECTO}}'}
    for marker, value in replacements.items():
        if marker in special:
            continue
        replace_marker_text(doc, marker, value)

    replace_marker_text_formatted(doc, '{{RESUMEN_NOMBRE_PROYECTO}}', project_name, bold=True, color_hex=BLACK, align='center')
    replace_marker_text_formatted(doc, '{{RESUMEN_RIESGO_GLOBAL}}', global_level, bold=True, color_hex=risk_color_hex(global_level), font_size=13, align='center')

    return replacements


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--case_id', default='')
    parser.add_argument('--data_path', default='')
    parser.add_argument('--template_path', default='/data/templates/NuevaPlantillaInforme_MARKERS.docx')
    parser.add_argument('--output_path', default='')
    args = parser.parse_args()

    case_id = args.case_id
    if not case_id and args.data_path:
        parts = Path(args.data_path).parts
        if 'cases' in parts:
            case_id = parts[parts.index('cases') + 1]
    if not case_id:
        raise Exception('Debe indicarse --case_id o --data_path dentro de /data/cases/<case_id>.')

    base = Path('/data/cases') / case_id
    report_dir = base / '04_report'
    output_dir = base / '05_report'
    output_dir.mkdir(parents=True, exist_ok=True)
    assets_dir = report_dir / 'assets'
    assets_dir.mkdir(parents=True, exist_ok=True)

    data_path = Path(args.data_path) if args.data_path else report_dir / 'report_render_data.json'
    template_path = Path(args.template_path)
    output_path = Path(args.output_path) if args.output_path else output_dir / f'Informe_Ciberseguridad_{case_id}.docx'

    data = load_json(data_path)

    # La plantilla incluye objetos OLE/Excel en cuadros de texto. Word puede
    # mostrar el aviso "No puede poner un objeto de dibujo dentro de un cuadro
    # de texto..." aunque el DOCX se genere correctamente. Por eso saneamos
    # una copia temporal de la plantilla antes de abrirla con python-docx.
    sanitized_template_path = output_dir / f'_sanitized_template_{case_id}.docx'
    shutil.copy2(template_path, sanitized_template_path)
    remove_ole_objects_inside_textboxes(sanitized_template_path)

    doc = Document(sanitized_template_path)

    xml_replacements = patch_basic_markers(doc, data)

    # Excel reference visible with hyperlink
    p = find_paragraph_with_marker(doc, '{{INSERT_CUESTIONARIO_OBJETO}}')
    if p:
        insert_excel_reference(doc, p, data, output_dir)

    # Incumplimientos grouped table
    p = find_paragraph_with_marker(doc, '{{INSERT_TABLA_INCUMPLIMIENTOS}}')
    if p:
        create_incumplimientos_table(doc, p, data.get('risk_groups', []))

    # Pie chart like the visual reference, without the extra left 'Riesgos totales' column.
    p = find_paragraph_with_marker(doc, '{{INSERT_GRAFICO_RIESGOS}}')
    if p:
        chart_path = assets_dir / 'riesgos_totales.png'
        make_risk_pie_chart(data.get('executive_visual', {}).get('risk_totals', {}), chart_path)
        insert_risk_chart_at_marker(doc, p, chart_path, width_inches=4.8)

    # Results sections
    p = find_paragraph_with_marker(doc, '{{INSERT_RESULTADOS_ANALISIS}}')
    if p:
        create_results_sections(doc, p, data.get('risk_groups', []))

    # Technology debt
    p = find_paragraph_with_marker(doc, '{{INSERT_DEUDA_TECNOLOGICA}}')
    if p:
        create_debt_table(doc, p, data.get('deuda_tecnologica', []))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    # Evita aviso de Word por objetos OLE/dibujos dentro de cuadros de texto de la plantilla.
    remove_ole_objects_inside_textboxes(output_path)

    # Patch markers in shapes/textboxes, especially cover text.
    replace_all_text_markers_in_xml(output_path, xml_replacements)
    postprocess_dynamic_report_xml(output_path, data)
    set_update_fields_on_open(output_path)

    # Save small generation metadata
    metadata = {
        'case_id': case_id,
        'status': 'DOCX_GENERATED',
        'created_at': datetime.now().isoformat(timespec='seconds'),
        'docx_path': str(output_path),
        'template_path': str(template_path),
        'data_path': str(data_path)
    }
    save_json(output_path.parent / 'report_generation_status.json', metadata)
    print(json.dumps(metadata, ensure_ascii=False))


if __name__ == '__main__':
    main()
