#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
rag_qdrant.py

Script para indexar documentos RAG en Qdrant usando embeddings locales de Ollama,
y para consultar posteriormente los fragmentos relevantes.

Diseñado para el flujo n8n local/self-hosted del TFM:
- n8n dentro de Docker
- Qdrant en http://qdrant:6333
- Ollama en http://ollama:11434
- Documentos RAG en /data/rag_docs

Comandos principales:
  index   -> indexa documentos en Qdrant
  query   -> consulta Qdrant con una pregunta
  health  -> comprueba Ollama y Qdrant
  list    -> lista colecciones de Qdrant
  reset   -> borra una colección
"""

import argparse
import hashlib
import json
import re
import sys
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib import request, error

try:
    from docx import Document
except Exception:
    Document = None


SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".docx"}
DEFAULT_COLLECTION = "tfm_rag_riesgos"
DEFAULT_QDRANT_URL = "http://qdrant:6333"
DEFAULT_OLLAMA_URL = "http://ollama:11434"
DEFAULT_EMBEDDING_MODEL = "nomic-embed-text"


# ============================================================
# UTILIDADES GENERALES
# ============================================================

def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ")
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def infer_source_type(path: Path) -> str:
    name = path.name.lower()
    if "riesgo" in name or "catalog" in name or "risk" in name:
        return "catalogo_riesgos"
    if "regla" in name or "business" in name:
        return "reglas_negocio"
    if "redaccion" in name or "guia" in name or "informe" in name:
        return "guia_redaccion_informe"
    if "plantilla" in name or "template" in name:
        return "plantilla_informe"
    return "documento_rag"


def safe_collection_name(name: str) -> str:
    name = re.sub(r"[^a-zA-Z0-9_\-]", "_", name.strip())
    return name or DEFAULT_COLLECTION


def make_point_id(source_sha: str, chunk_index: int) -> str:
    """Qdrant acepta UUIDs como IDs. Usamos UUID determinista para evitar duplicados."""
    return str(uuid.uuid5(uuid.NAMESPACE_URL, f"{source_sha}:{chunk_index}"))


# ============================================================
# HTTP JSON SIN DEPENDENCIAS EXTERNAS
# ============================================================

def http_json(method: str, url: str, payload: dict[str, Any] | None = None, timeout: int = 120) -> dict[str, Any]:
    data = None
    headers = {"Content-Type": "application/json"}

    if payload is not None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")

    req = request.Request(url, data=data, headers=headers, method=method.upper())

    try:
        with request.urlopen(req, timeout=timeout) as resp:
            raw = resp.read().decode("utf-8")
            if not raw:
                return {}
            return json.loads(raw)
    except error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"HTTP {exc.code} en {url}: {body}") from exc
    except error.URLError as exc:
        raise RuntimeError(f"No se pudo conectar con {url}: {exc}") from exc


# ============================================================
# OLLAMA EMBEDDINGS
# ============================================================

def ollama_embed(text: str, ollama_url: str, model: str) -> list[float]:
    text = normalize_text(text)
    if not text:
        raise ValueError("No se puede generar embedding de texto vacío")

    base = ollama_url.rstrip("/")

    # API clásica: /api/embeddings
    try:
        res = http_json(
            "POST",
            f"{base}/api/embeddings",
            {"model": model, "prompt": text},
            timeout=180,
        )
        emb = res.get("embedding")
        if isinstance(emb, list) and emb:
            return [float(x) for x in emb]
    except Exception:
        pass

    # API nueva: /api/embed
    res = http_json(
        "POST",
        f"{base}/api/embed",
        {"model": model, "input": text},
        timeout=180,
    )
    embeddings = res.get("embeddings")
    if isinstance(embeddings, list) and embeddings:
        emb = embeddings[0]
        return [float(x) for x in emb]

    raise RuntimeError(f"Ollama no devolvió embedding válido con el modelo {model}")


# ============================================================
# QDRANT REST
# ============================================================

def qdrant_collection_exists(qdrant_url: str, collection: str) -> bool:
    base = qdrant_url.rstrip("/")
    try:
        http_json("GET", f"{base}/collections/{collection}", timeout=30)
        return True
    except Exception:
        return False


def qdrant_create_collection(qdrant_url: str, collection: str, vector_size: int, distance: str = "Cosine") -> None:
    base = qdrant_url.rstrip("/")
    payload = {
        "vectors": {
            "size": vector_size,
            "distance": distance,
        }
    }
    http_json("PUT", f"{base}/collections/{collection}", payload, timeout=120)


def qdrant_delete_collection(qdrant_url: str, collection: str) -> None:
    base = qdrant_url.rstrip("/")
    try:
        http_json("DELETE", f"{base}/collections/{collection}", timeout=120)
    except Exception:
        pass


def qdrant_upsert_points(qdrant_url: str, collection: str, points: list[dict[str, Any]]) -> None:
    if not points:
        return
    base = qdrant_url.rstrip("/")
    payload = {"points": points}
    http_json("PUT", f"{base}/collections/{collection}/points?wait=true", payload, timeout=240)


def qdrant_search(
    qdrant_url: str,
    collection: str,
    vector: list[float],
    top_k: int,
    score_threshold: float | None = None,
) -> list[dict[str, Any]]:
    base = qdrant_url.rstrip("/")
    payload: dict[str, Any] = {
        "vector": vector,
        "limit": top_k,
        "with_payload": True,
        "with_vector": False,
    }
    if score_threshold is not None:
        payload["score_threshold"] = score_threshold

    res = http_json("POST", f"{base}/collections/{collection}/points/search", payload, timeout=120)
    return res.get("result", [])


def qdrant_list_collections(qdrant_url: str) -> dict[str, Any]:
    base = qdrant_url.rstrip("/")
    return http_json("GET", f"{base}/collections", timeout=30)


# ============================================================
# LECTURA DE DOCUMENTOS
# ============================================================

def read_txt_or_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def flatten_json(obj: Any, prefix: str = "") -> list[str]:
    lines = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            new_prefix = f"{prefix}.{k}" if prefix else str(k)
            lines.extend(flatten_json(v, new_prefix))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            new_prefix = f"{prefix}[{i}]"
            lines.extend(flatten_json(v, new_prefix))
    else:
        lines.append(f"{prefix}: {obj}")
    return lines


def read_json_file(path: Path) -> str:
    obj = json.loads(path.read_text(encoding="utf-8"))
    return "\n".join(flatten_json(obj))


def read_docx(path: Path) -> str:
    if Document is None:
        raise RuntimeError("python-docx no está instalado. Instala python-docx o usa .txt/.md")

    doc = Document(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        txt = normalize_text(p.text)
        if txt:
            parts.append(txt)

    for t_index, table in enumerate(doc.tables, start=1):
        parts.append(f"Tabla {t_index}:")
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def read_document(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md"}:
        return read_txt_or_md(path)
    if ext == ".json":
        return read_json_file(path)
    if ext == ".docx":
        return read_docx(path)
    raise ValueError(f"Extensión no soportada para RAG: {path.suffix}")


def iter_documents(docs_dir: Path) -> list[Path]:
    if not docs_dir.exists():
        raise FileNotFoundError(f"No existe la carpeta de documentos RAG: {docs_dir}")

    files = []
    for path in sorted(docs_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            if path.name.startswith("~$"):
                continue
            files.append(path)

    return files


# ============================================================
# CHUNKING
# ============================================================

def split_long_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    chunks = []
    start = 0
    text = normalize_text(text)
    n = len(text)

    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - chunk_overlap)

    return chunks


def chunk_document(text: str, chunk_size: int = 1200, chunk_overlap: int = 200) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            if current.strip():
                chunks.append(current.strip())
                current = ""
            chunks.extend(split_long_text(paragraph, chunk_size, chunk_overlap))
            continue

        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current.strip():
                chunks.append(current.strip())
            current = paragraph

    if current.strip():
        chunks.append(current.strip())

    # Añadir overlap textual entre chunks para mantener continuidad.
    if chunk_overlap > 0 and len(chunks) > 1:
        overlapped = []
        previous_tail = ""
        for idx, chunk in enumerate(chunks):
            if idx == 0:
                overlapped.append(chunk)
            else:
                merged = f"{previous_tail}\n\n{chunk}".strip()
                overlapped.append(merged)
            previous_tail = chunk[-chunk_overlap:]
        chunks = overlapped

    return chunks


# ============================================================
# INDEXACIÓN Y CONSULTA
# ============================================================

def build_points_from_docs(
    docs_dir: Path,
    qdrant_url: str,
    ollama_url: str,
    embedding_model: str,
    chunk_size: int,
    chunk_overlap: int,
) -> tuple[list[dict[str, Any]], int, list[dict[str, Any]]]:
    docs = iter_documents(docs_dir)
    points: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    vector_size = 0

    for doc_path in docs:
        try:
            raw_text = read_document(doc_path)
            chunks = chunk_document(raw_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            source_sha = file_sha256(doc_path)
            source_type = infer_source_type(doc_path)

            for idx, chunk in enumerate(chunks):
                embedding = ollama_embed(chunk, ollama_url=ollama_url, model=embedding_model)
                if vector_size == 0:
                    vector_size = len(embedding)

                payload = {
                    "text": chunk,
                    "source_file": doc_path.name,
                    "source_path": str(doc_path),
                    "source_type": source_type,
                    "source_sha256": source_sha,
                    "chunk_index": idx,
                    "chunk_count": len(chunks),
                    "text_preview": chunk[:300],
                    "indexed_at": now_iso(),
                }

                points.append(
                    {
                        "id": make_point_id(source_sha, idx),
                        "vector": embedding,
                        "payload": payload,
                    }
                )

        except Exception as exc:
            errors.append(
                {
                    "file": str(doc_path),
                    "error_type": type(exc).__name__,
                    "error": str(exc),
                }
            )

    return points, vector_size, errors


def index_documents(args) -> dict[str, Any]:
    collection = safe_collection_name(args.collection)
    docs_dir = Path(args.docs_dir)

    points, vector_size, errors = build_points_from_docs(
        docs_dir=docs_dir,
        qdrant_url=args.qdrant_url,
        ollama_url=args.ollama_url,
        embedding_model=args.embedding_model,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
    )

    if not points:
        return {
            "ok": False,
            "command": "index",
            "collection": collection,
            "message": "No se generó ningún punto para indexar",
            "errors": errors,
        }

    if args.recreate:
        qdrant_delete_collection(args.qdrant_url, collection)
        time.sleep(0.5)

    if not qdrant_collection_exists(args.qdrant_url, collection):
        qdrant_create_collection(args.qdrant_url, collection, vector_size=vector_size)

    batch_size = args.batch_size
    indexed = 0

    for start in range(0, len(points), batch_size):
        batch = points[start:start + batch_size]
        qdrant_upsert_points(args.qdrant_url, collection, batch)
        indexed += len(batch)

    docs = iter_documents(docs_dir)
    return {
        "ok": True,
        "command": "index",
        "collection": collection,
        "documents_found": len(docs),
        "points_indexed": indexed,
        "vector_size": vector_size,
        "embedding_model": args.embedding_model,
        "qdrant_url": args.qdrant_url,
        "ollama_url": args.ollama_url,
        "errors": errors,
    }


def query_documents(args) -> dict[str, Any]:
    collection = safe_collection_name(args.collection)

    if not qdrant_collection_exists(args.qdrant_url, collection):
        return {
            "ok": False,
            "command": "query",
            "collection": collection,
            "message": "La colección no existe. Ejecuta primero el comando index.",
        }

    embedding = ollama_embed(args.text, ollama_url=args.ollama_url, model=args.embedding_model)
    hits = qdrant_search(
        qdrant_url=args.qdrant_url,
        collection=collection,
        vector=embedding,
        top_k=args.top_k,
        score_threshold=args.score_threshold,
    )

    formatted_hits = []
    for hit in hits:
        payload = hit.get("payload", {}) or {}
        formatted_hits.append(
            {
                "score": hit.get("score"),
                "source_file": payload.get("source_file"),
                "source_type": payload.get("source_type"),
                "chunk_index": payload.get("chunk_index"),
                "text": payload.get("text"),
                "text_preview": payload.get("text_preview"),
            }
        )

    return {
        "ok": True,
        "command": "query",
        "collection": collection,
        "query": args.text,
        "top_k": args.top_k,
        "hits_count": len(formatted_hits),
        "hits": formatted_hits,
    }


def health(args) -> dict[str, Any]:
    result: dict[str, Any] = {
        "ok": True,
        "command": "health",
        "qdrant": None,
        "ollama": None,
    }

    try:
        result["qdrant"] = qdrant_list_collections(args.qdrant_url)
    except Exception as exc:
        result["ok"] = False
        result["qdrant"] = {
            "ok": False,
            "error": str(exc),
        }

    try:
        base = args.ollama_url.rstrip("/")
        result["ollama"] = http_json("GET", f"{base}/api/tags", timeout=30)
    except Exception as exc:
        result["ok"] = False
        result["ollama"] = {
            "ok": False,
            "error": str(exc),
        }

    return result


def list_collections(args) -> dict[str, Any]:
    return {
        "ok": True,
        "command": "list",
        "qdrant_url": args.qdrant_url,
        "collections": qdrant_list_collections(args.qdrant_url),
    }


def reset_collection(args) -> dict[str, Any]:
    collection = safe_collection_name(args.collection)
    qdrant_delete_collection(args.qdrant_url, collection)
    return {
        "ok": True,
        "command": "reset",
        "collection": collection,
        "message": "Colección eliminada si existía",
    }


# ============================================================
# MAIN
# ============================================================

def print_result(result: dict[str, Any]) -> None:
    print(json.dumps(result, ensure_ascii=False, indent=2))


def main():
    parser = argparse.ArgumentParser(
        description="RAG local con Qdrant + Ollama para el TFM"
    )

    subparsers = parser.add_subparsers(dest="command", required=True)

    # health
    p_health = subparsers.add_parser("health", help="Comprueba conexión con Qdrant y Ollama")
    p_health.add_argument("--qdrant-url", default=DEFAULT_QDRANT_URL)
    p_health.add_argument("--ollama-url", default=DEFAULT_OLLAMA_URL)
    p_health.set_defaults(func=health)

    # list
    p_list = subparsers.add_parser("list", help="Lista colecciones de Qdrant")
    p_list.add_argument("--qdrant-url", default=DEFAULT_QDRANT_URL)
    p_list.set_defaults(func=list_collections)

    # reset
    p_reset = subparsers.add_parser("reset", help="Elimina una colección de Qdrant")
    p_reset.add_argument("--qdrant-url", default=DEFAULT_QDRANT_URL)
    p_reset.add_argument("--collection", default=DEFAULT_COLLECTION)
    p_reset.set_defaults(func=reset_collection)

    # index
    p_index = subparsers.add_parser("index", help="Indexa documentos RAG en Qdrant")
    p_index.add_argument("--docs-dir", required=True, type=Path)
    p_index.add_argument("--collection", default=DEFAULT_COLLECTION)
    p_index.add_argument("--qdrant-url", default=DEFAULT_QDRANT_URL)
    p_index.add_argument("--ollama-url", default=DEFAULT_OLLAMA_URL)
    p_index.add_argument("--embedding-model", default=DEFAULT_EMBEDDING_MODEL)
    p_index.add_argument("--chunk-size", type=int, default=1200)
    p_index.add_argument("--chunk-overlap", type=int, default=200)
    p_index.add_argument("--batch-size", type=int, default=32)
    p_index.add_argument("--recreate", action="store_true")
    p_index.set_defaults(func=index_documents)

    # query
    p_query = subparsers.add_parser("query", help="Consulta documentos RAG en Qdrant")
    p_query.add_argument("--text", required=True)
    p_query.add_argument("--collection", default=DEFAULT_COLLECTION)
    p_query.add_argument("--qdrant-url", default=DEFAULT_QDRANT_URL)
    p_query.add_argument("--ollama-url", default=DEFAULT_OLLAMA_URL)
    p_query.add_argument("--embedding-model", default=DEFAULT_EMBEDDING_MODEL)
    p_query.add_argument("--top-k", type=int, default=5)
    p_query.add_argument("--score-threshold", type=float, required=False)
    p_query.set_defaults(func=query_documents)

    args = parser.parse_args()

    try:
        result = args.func(args)
        print_result(result)
    except Exception as exc:
        error_result = {
            "ok": False,
            "command": args.command,
            "error_type": type(exc).__name__,
            "error": str(exc),
        }
        print_result(error_result)
        sys.exit(1)


if __name__ == "__main__":
    main()
